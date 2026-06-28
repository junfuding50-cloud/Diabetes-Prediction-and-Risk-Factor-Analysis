# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

def set_east_asia(run, font='SimSun'):
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)

def sf(run, size=12, bold=False, color=None):
    set_east_asia(run)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def h(text, level=1):
    p = doc.add_heading(level=level)
    p.clear()
    r = p.add_run(text)
    sf(r, size=14 if level==1 else 13, bold=True)

def para(text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    sf(r)

def img_placeholder(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'【请在此处插入图片：{text}】')
    sf(r, size=11, color=(150,150,150))

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i,h_ in enumerate(headers):
        r = t.rows[0].cells[i].paragraphs[0].add_run(h_)
        sf(r, bold=True, size=11)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            r = t.rows[ri+1].cells[ci].paragraphs[0].add_run(str(val))
            sf(r, size=11)
    doc.add_paragraph()

# 标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('基于机器学习的糖尿病预测与风险因素分析')
sf(r, size=18, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('人工智能导论（医学5年制）实践课题报告')
sf(r, size=13)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('姓名/学号/专业/年级：___________________________')
sf(r)
doc.add_page_break()

# 摘要
h('摘要')
para('本研究基于皮马印第安人糖尿病数据集（768条记录），使用Python/Pandas完成数据清洗，并借助Orange Data Mining工具构建可视化机器学习工作流，对糖尿病患病风险进行预测分析并识别关键风险因素。研究发现，在全特征模型中梯度提升算法表现最优，AUC达0.941，准确率87.7%。进一步分析发现，血糖与胰岛素指标存在反向因果问题，将其纳入模型会引入数据泄露风险；排除后重新建模，随机森林成为最优模型（AUC=0.845），梯度提升AUC降至0.819（降幅0.122），证实了上述特征对模型的过度贡献。在剩余特征中，年龄、皮褶厚度和BMI为具有真实预防意义的关键风险因素。本研究揭示了医疗数据挖掘中数据泄露问题对模型可信度的影响，并为糖尿病早期筛查的特征选择提供了参考依据。', indent=True)
p = doc.add_paragraph()
r = p.add_run('关键词：'); sf(r, bold=True)
r = p.add_run('糖尿病预测；机器学习；风险因素分析；数据泄露；Orange Data Mining'); sf(r)
doc.add_paragraph()

# 一、引言
h('一、引言/背景')
para('糖尿病是全球发病率持续上升的慢性代谢疾病，2型糖尿病尤为普遍。皮马印第安人因遗传背景特殊，2型糖尿病患病率显著高于一般人群，其体检数据为研究糖尿病风险因素提供了宝贵样本。', indent=True)
para('本研究目标：①基于多项体检指标构建糖尿病预测模型，对比多种机器学习算法的预测性能；②识别与糖尿病患病关联最强的关键风险因素；③探讨医疗数据挖掘中数据泄露问题对模型解释性与可信度的影响。', indent=True)

# 二、材料与方法
h('二、材料与方法')
h('2.1 数据集', level=2)
para('本研究使用Pima Indians Diabetes Database（UCI机器学习仓库）。数据集包含768条记录，8项体检特征及1项患病标签（Outcome）。样本中268例患病（34.9%），500例未患病（65.1%），存在轻度类别不平衡，故评估指标除准确率外还参考AUC和F1值。', indent=True)
h('2.2 数据清洗（Python/Pandas）', level=2)
table(
    ['特征','含义'],
    [['Pregnancies','怀孕次数'],['Glucose','血糖浓度'],['BloodPressure','舒张压（mmHg）'],
     ['SkinThickness','三头肌皮褶厚度（mm）'],['Insulin','2小时血清胰岛素（μU/ml）'],
     ['BMI','身体质量指数'],['DiabetesPedigreeFunction','糖尿病家族遗传函数'],
     ['Age','年龄（岁）'],['Outcome','是否患糖尿病（1=是，0=否）']]
)

h('2.2 数据清洗', level=2)
para('数据清洗在Python中使用Pandas库完成。数据集中Glucose、BloodPressure、SkinThickness、Insulin、BMI字段存在大量数值0，而这些生理指标在医学上不可能为0，实为缺失值被填充为0（Insulin缺失率48.7%，SkinThickness缺失率29.6%）。处理策略：使用Pandas的.mask()方法将0值替换为缺失值（NaN），再按患病分组（Outcome=0/1）用各组中位数填充，比全局填充更合理。清洗后数据保存为clean_data.csv，导入Orange进行后续分析。', indent=True)
img_placeholder('插图1：Orange工作流整体截图（画布全貌）')

h('2.3 分析工具与模型', level=2)
para('工具：Orange Data Mining（可视化拖拽式数据挖掘平台）', indent=True)
para('分析组件：Distributions（特征分布）、Rank（特征重要性）、Box Plot（箱线图）、Scatter Plot（散点图）')
para('预测模型：决策树（Decision Tree）、随机森林（Random Forest）、梯度提升（Gradient Boosting）')
para('评估方法：Test and Score，10折交叉验证，评估指标：AUC、准确率（CA）、F1、召回率（Recall）、MCC')

# 三、结果展示
h('三、结果展示')
h('3.1 特征分布与初步探索', level=2)
img_placeholder('插图2：Distributions截图（Glucose或BMI，Split by Outcome）')
img_placeholder('插图3：Box Plot截图（Age或BMI，Subgroups=Outcome）')
para('各特征分布分析显示，患病人群在血糖、BMI、年龄等指标上的分布明显高于未患病人群，初步表明上述指标与糖尿病存在关联。', indent=True)

h('3.2 特征重要性排序（Rank统计分析）', level=2)
img_placeholder('插图4：Rank特征重要性排序截图')
para('Rank分析基于信息增益和基尼系数对各特征区分患病的能力进行量化排序。在全部8个特征中，Glucose和Insulin排名最前，但其存在反向因果问题（详见3.4节）。排除上述两项后，年龄（Age）、皮褶厚度（SkinThickness）、BMI重要性评分居前，是真正值得关注的风险因素。需说明的是，SkinThickness原始缺失率约29.6%，其评分需审慎解读；BMI与SkinThickness同属肥胖相关指标，共同指向肥胖是糖尿病的重要可干预风险因素。', indent=True)

h('3.3 机器学习模型特征重要性（Python/Sklearn）', level=2)
para('为弥补Orange在集成模型特征重要性输出方面的局限，本研究使用Python/Sklearn对随机森林和梯度提升模型进行特征重要性分析。随机森林的特征重要性基于各特征在所有决策树中降低基尼不纯度的平均贡献；梯度提升则基于各特征在迭代过程中减少损失函数的累积贡献。', indent=True)
img_placeholder('插图5：Python生成的特征重要性对比图（feature_importance.png，左：随机森林，右：梯度提升）')
para('如图所示，两模型均显示Glucose和Insulin重要性最高，从模型角度印证了其与患病标签的强相关性，同时也进一步证实了将其纳入预测的数据泄露风险。此外，SkinThickness在两模型间重要性差异较大：一方面，SkinThickness与BMI高度相关，梯度提升的序列建树方式使两者产生竞争，重要性集中于其中一个；另一方面，29.6%的高缺失率导致大量样本填充为相同的中位数，降低了梯度提升残差修正阶段的可用信息量。排除Glucose和Insulin后，两模型均指向年龄、BMI为最可靠的风险因素，与Rank统计分析结论一致，互相印证。', indent=True)

h('3.4 数据泄露分析：两次建模对比', level=2)
para('本研究设计消融实验，对比包含与排除高泄露风险特征的模型性能差异。', indent=True)
para('模型A——全特征（防止数据泄露前）：')
img_placeholder('插图5：防止数据泄露前的模型对比截图')
table(
    ['模型','AUC','CA','F1','Recall'],
    [['决策树','0.831','0.852','0.851','0.852'],
     ['随机森林','0.938','0.874','0.874','0.874'],
     ['梯度提升','0.941','0.877','0.877','0.877']]
)
para('模型B——排除血糖/胰岛素后（防止数据泄露后）：')
img_placeholder('插图6：防止数据泄露后的模型对比截图')
table(
    ['模型','AUC','CA','F1','Recall'],
    [['决策树','0.651','0.714','0.713','0.714'],
     ['随机森林','0.845','0.774','0.771','0.774'],
     ['梯度提升','0.819','0.754','0.753','0.754']]
)
para('排除Glucose和Insulin后，梯度提升AUC由0.941下降至0.819（降幅0.122），准确率由87.7%降至75.4%（降幅12.3个百分点），证实了这两个特征对模型的显著贡献。', indent=True)

# 四、讨论
h('四、讨论分析')
h('4.1 模型性能分析', level=2)
para('三种模型中，梯度提升综合表现最优（全特征AUC=0.941），其次为随机森林（AUC=0.938），单棵决策树最弱（AUC=0.831）。排除血糖和胰岛素后，随机森林成为最优模型（AUC=0.845），梯度提升下降至0.819，说明梯度提升对泄露特征的依赖程度更高。集成模型整体明显优于单一决策树，原因在于集成方法通过多个弱学习器的组合降低了方差，对特征噪声较多的医疗数据具有更强鲁棒性。在医疗筛查场景中，漏诊的危害远大于误报；梯度提升全特征召回率达0.877，具有较好的临床筛查价值。', indent=True)

h('4.2 数据泄露的影响', level=2)
para('血糖与胰岛素是临床诊断糖尿病的直接依据，二者本质上是糖尿病的表现（果）而非原因（因）。将其纳入预测模型，存在因果方向倒置的问题，在早期筛查场景下意义有限。消融实验直观证明了这一点：去除这两个特征后，AUC下降约0.122，说明全特征模型的高准确率有相当部分依赖于对症状型指标的拟合。此外，Insulin缺失率高达48.7%，大量中位数填充引入噪声，也是其被排除的合理依据。', indent=True)

h('4.3 关键风险因素的医学意义', level=2)
para('年龄（Age）：随年龄增长，胰岛素敏感性下降、β细胞功能减退，患病风险递增。年龄不可干预，但提示应对中老年人群进行重点筛查。', indent=True)
para('BMI / 皮褶厚度：两者均反映肥胖程度。肥胖是2型糖尿病最重要的可干预风险因素，通过控制体重和增加运动可有效降低患病风险。皮褶厚度缺失率较高，结论需审慎看待。', indent=True)

h('4.4 AI使用声明', level=2)
para('本报告数据分析与模型训练使用Orange Data Mining完成。报告撰写过程中借助AI工具（Claude）辅助生成文字表达与结构组织；分析思路、实验设计及结论判断均由作者独立完成，AI工具用于文字润色与格式整理。', indent=True)

# 五、结论
h('五、结论与展望')
para('主要结论：①梯度提升算法在全特征下综合性能最优（AUC=0.941），排除泄露特征后随机森林成为最优模型（AUC=0.845）；②血糖和胰岛素存在反向因果问题，消融实验定量证明其对模型的过度贡献（梯度提升AUC降幅0.122）；③排除泄露特征后，年龄、BMI、皮褶厚度是具有真实预防意义的关键风险因素，肥胖相关指标因其可干预性尤为重要；④数据清洗（Python/Pandas去除伪零值）是保证模型可靠性的重要前提。', indent=True)
para('局限性：样本量较小（768条），部分特征缺失率高，研究对象限定为特定族裔女性，结论外推需谨慎。', indent=True)
para('未来展望：可引入更大规模多来源数据集，采用多重插补等更精细的缺失值处理方法，并探索深度学习在医疗预测中的应用潜力。', indent=True)

# 六、参考文献
h('六、参考文献')
for ref in [
    '[1] Smith J W, et al. Using the ADAP learning algorithm to forecast the onset of diabetes mellitus[C]. SCAMC, 1988: 261-265.',
    '[2] UCI Machine Learning Repository. Pima Indians Diabetes Database[EB/OL]. https://archive.ics.uci.edu/ml/datasets/diabetes.',
    '[3] Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.',
    '[4] Friedman J H. Greedy function approximation: a gradient boosting machine[J]. Annals of Statistics, 2001, 29(5): 1189-1232.',
    '[5] Demšar J, et al. Orange: data mining toolbox in Python[J]. JMLR, 2013, 14(1): 2349-2353.',
]:
    para(ref)

doc.save('D:/final term ai work/report_v3.docx')
print('done')
